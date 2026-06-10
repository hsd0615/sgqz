#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成武将10/50/100/200级数值表格
"""

import os
import sys
import xml.etree.ElementTree as ET

# Paths
BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
GENERAL_XML = os.path.join(BASE, "staticgeneral.xml")
XISHU_XML = os.path.join(BASE, "staticxishu.xml")

# Type mapping
TYPE_NAMES = {
    0: "投石车", 1: "弓兵", 2: "飞刀兵", 3: "朴刀兵",
    4: "斧兵", 5: "锤兵", 6: "武斗兵", 7: "长枪兵",
    8: "藤甲兵", 9: "骑兵", 10: "弯刀兵", 11: "箭塔兵",
    12: "强弓兵", 13: "BOSS", 20: "君主"
}

# Title mapping
TITLE_NAMES = {
    0: "超级武将",
    1: "一流武将",
    2: "二流武将",
    3: "三流武将"
}

# Feature / 五行属性
FEATURE_NAMES = {
    0: "无",
    1: "金",
    2: "木",
    3: "水",
    4: "火",
}

def parse_xml(path):
    """Parse XML file, return root element"""
    with open(path, 'r', encoding='utf-8') as f:
        content = f.read()
    # Fix unclosed tags like <jinbi/> <coint/>
    root = ET.fromstring(content)
    return root

def parse_generals():
    """Parse all generals from XML"""
    root = parse_xml(GENERAL_XML)
    generals = []
    for record in root.findall('RECORD'):
        g = {}
        g['code'] = (record.find('code').text or '').strip() if record.find('code') is not None else ''
        g['name'] = (record.find('name').text or '').strip() if record.find('name') is not None else ''
        g['type'] = int((record.find('type').text or '0').strip()) if record.find('type') is not None else 0
        g['title'] = int((record.find('title').text or '3').strip()) if record.find('title') is not None else 3
        g['hp'] = int((record.find('hp').text or '0').strip()) if record.find('hp') is not None else 0
        g['attack'] = int((record.find('attack').text or '0').strip()) if record.find('attack') is not None else 0
        g['defense'] = int((record.find('defense').text or '0').strip()) if record.find('defense') is not None else 0

        # Parse kezhi
        kezhi_el = record.find('kezhi')
        g['kezhi'] = (kezhi_el.text or '').strip() if kezhi_el is not None and kezhi_el.text else ''

        g['skin'] = (record.find('skin').text or '').strip() if record.find('skin') is not None else ''

        # Recruitment info
        g['recruitLevel'] = int((record.find('recruitLevel').text or '0').strip()) if record.find('recruitLevel') is not None else 0
        g['money'] = int((record.find('money').text or '0').strip()) if record.find('money') is not None else 0

        jinbi_el = record.find('jinbi')
        g['jinbi'] = int((jinbi_el.text or '0').strip()) if jinbi_el is not None and jinbi_el.text else 0

        dianka_el = record.find('dianka')
        g['dianka'] = int((dianka_el.text or '0').strip()) if dianka_el is not None and dianka_el.text else 0

        coint_el = record.find('coint')
        g['coint'] = int((coint_el.text or '0').strip()) if coint_el is not None and coint_el.text else 0

        price_el = record.find('price')
        g['price'] = int((price_el.text or '0').strip()) if price_el is not None and price_el.text else 0

        proto_el = record.find('proto')
        g['proto'] = (proto_el.text or '').strip() if proto_el is not None and proto_el.text else ''

        generals.append(g)
    return generals

def parse_xishu():
    """Parse growth coefficients"""
    root = parse_xishu(XISHU_XML)
    xishu = {}
    for record in root.findall('RECORD'):
        code = (record.find('code').text or '').strip() if record.find('code') is not None else ''
        hp = float((record.find('hp').text or '1').strip()) if record.find('hp') is not None else 1.0
        attack = float((record.find('attack').text or '1').strip()) if record.find('attack') is not None else 1.0
        defence = float((record.find('defence').text or '1').strip()) if record.find('defence') is not None else 1.0
        xishu[code] = {'hp': hp, 'attack': attack, 'defence': defence}
    return xishu

def calc_base_hp(gtype, level, base_hp, title, xishu):
    """Calculate base HP using game formula"""
    code = f"{gtype}_{title}"
    xs = xishu.get(code, {'hp': 1.0})
    xh = xs['hp']
    part1 = (level - 1) * 50 * xh
    part2 = base_hp * (1 + level * 0.03)
    return int(part1 + part2)

def calc_base_attack(gtype, level, base_attack, title, xishu):
    """Calculate base attack using game formula"""
    code = f"{gtype}_{title}"
    xs = xishu.get(code, {'attack': 1.0})
    xa = xs['attack']
    part1 = (level - 1) * 30 * xa
    part2 = base_attack * (1 + level * 0.03)
    return int(part1 + part2)

def calc_base_defense(gtype, level, base_defense, title, xishu):
    """Calculate base defense using game formula"""
    code = f"{gtype}_{title}"
    xs = xishu.get(code, {'defence': 1.0})
    xd = xs['defence']
    part1 = (level - 1) * 30 * xd
    part2 = base_defense * (1 + level * 0.03)
    return int(part1 + part2)

def get_acquisition(g):
    """Determine acquisition method"""
    methods = []
    rl = g['recruitLevel']
    price = g['price']
    money = g['money']
    jinbi = g['jinbi']
    dianka = g['dianka']
    coint = g['coint']

    # System default (recruitLevel == 1000)
    if rl == 1000:
        methods.append("系统默认")
    elif rl > 0:
        methods.append(f"招募(Lv{rl})")

    # Shop purchase
    if price > 0:
        methods.append(f"点卡购买({price}元)")
    elif price == -1 and rl <= 0:
        # Not recruitable normally, no price - might be unobtainable or event-based
        if rl == -1:
            methods.append("活动/特殊获取")
        elif rl == 0:
            methods.append("活动/特殊获取")

    # Cost info
    cost_parts = []
    if money > 0:
        cost_parts.append(f"{money}银两")
    if jinbi > 0:
        cost_parts.append(f"{jinbi}金币")
    if dianka > 0:
        cost_parts.append(f"{dianka}点卡")
    if coint > 0:
        cost_parts.append(f"{coint}元宝")

    if not methods:
        methods.append("活动/特殊获取")

    result = "/".join(methods)
    if cost_parts:
        result += f" [{', '.join(cost_parts)}]"
    return result

def apply_ju_yi_rename(generals):
    """将鞠义改为三流骑兵"""
    for g in generals:
        if g['name'] == '鞠义':
            g['name'] = '三流骑兵'
    return generals

def main():
    generals = parse_generals()
    xishu = parse_xishu()

    # Apply 鞠义 rename
    generals = apply_ju_yi_rename(generals)

    levels = [10, 50, 100, 200]
    desktop = os.path.join(os.path.expanduser("~"), "Desktop")

    # Output TXT
    txt_path = os.path.join(desktop, "武将数值表.txt")
    with open(txt_path, 'w', encoding='utf-8') as f:
        # Header
        header = f"{'武将名称':<10} {'类型':<8} {'品质':<10} {'获取方式':<30} "
        for lv in levels:
            header += f"{'Lv'+str(lv)+'生命':>8} {'Lv'+str(lv)+'攻击':>8} {'Lv'+str(lv)+'防御':>8}  "
        f.write(header + "\n")
        f.write("=" * (len(header) + 20) + "\n")

        for g in generals:
            gtype_name = TYPE_NAMES.get(g['type'], f"未知({g['type']})")
            title_name = TITLE_NAMES.get(g['title'], f"未知({g['title']})")
            acq = get_acquisition(g)

            name = g['name']
            line = f"{name:<10} {gtype_name:<8} {title_name:<10} {acq:<30} "

            for lv in levels:
                hp = calc_base_hp(g['type'], lv, g['hp'], g['title'], xishu)
                atk = calc_base_attack(g['type'], lv, g['attack'], g['title'], xishu)
                df = calc_base_defense(g['type'], lv, g['defense'], g['title'], xishu)
                line += f"{hp:>8} {atk:>8} {df:>8}  "

            f.write(line + "\n")

        # Summary
        f.write("\n" + "=" * 80 + "\n")
        f.write(f"总计 {len(generals)} 名武将\n")
        f.write(f"公式: HP = (等级-1)×50×系数 + 基础HP×(1+等级×0.03)\n")
        f.write(f"公式: 攻击 = (等级-1)×30×系数 + 基础攻击×(1+等级×0.03)\n")
        f.write(f"公式: 防御 = (等级-1)×30×系数 + 基础防御×(1+等级×0.03)\n")
        f.write(f"注: 鞠义已更名为三流骑兵\n")

    print(f"TXT file saved to: {txt_path}")

    # Try to generate DOCX if python-docx is available
    try:
        from docx import Document
        from docx.shared import Pt, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        doc = Document()
        doc.styles['Normal'].font.size = Pt(8)
        doc.styles['Normal'].font.name = 'SimSun'

        title = doc.add_paragraph('三国Q战4399版 - 武将数值表')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.size = Pt(14)
        title.runs[0].font.bold = True

        subtitle = doc.add_paragraph(f'总计 {len(generals)} 名武将 | 10级/50级/100级/200级 | 鞠义→三流骑兵')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Create table
        cols = 1 + 1 + 1 + 1 + len(levels) * 3  # 名称+类型+品质+获取+每级(HP/攻击/防御)
        table = doc.add_table(rows=1, cols=cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        hdr_cells = table.rows[0].cells
        headers = ['武将名称', '类型', '品质', '获取方式']
        for lv in levels:
            headers += [f'Lv{lv}\n生命', f'Lv{lv}\n攻击', f'Lv{lv}\n防御']

        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            for p in hdr_cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(7)
                    run.font.bold = True

        # Data rows
        for g in generals:
            row_cells = table.add_row().cells
            name = g['name']
            gtype_name = TYPE_NAMES.get(g['type'], f"未知({g['type']})")
            title_name = TITLE_NAMES.get(g['title'], f"未知({g['title']})")
            acq = get_acquisition(g)

            row_cells[0].text = name
            row_cells[1].text = gtype_name
            row_cells[2].text = title_name
            row_cells[3].text = acq

            col_idx = 4
            for lv in levels:
                hp = calc_base_hp(g['type'], lv, g['hp'], g['title'], xishu)
                atk = calc_base_attack(g['type'], lv, g['attack'], g['title'], xishu)
                df = calc_base_defense(g['type'], lv, g['defense'], g['title'], xishu)
                row_cells[col_idx].text = str(hp)
                row_cells[col_idx + 1].text = str(atk)
                row_cells[col_idx + 2].text = str(df)
                col_idx += 3

            # Set font size for all cells in this row
            for cell in row_cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(7)

        # Add formula note
        doc.add_paragraph()
        note = doc.add_paragraph()
        note.add_run('公式说明：').bold = True
        doc.add_paragraph('HP = (等级-1) × 50 × 成长系数 + 基础HP × (1 + 等级 × 0.03)')
        doc.add_paragraph('攻击 = (等级-1) × 30 × 成长系数 + 基础攻击 × (1 + 等级 × 0.03)')
        doc.add_paragraph('防御 = (等级-1) × 30 × 成长系数 + 基础防御 × (1 + 等级 × 0.03)')
        doc.add_paragraph('成长系数取自staticxishu.xml，由兵种类型+品质等级决定')
        doc.add_paragraph('注: 鞠义已更名为三流骑兵')

        docx_path = os.path.join(desktop, "武将数值表.docx")
        doc.save(docx_path)
        print(f"DOCX file saved to: {docx_path}")

    except ImportError:
        print("python-docx not available, skipping DOCX generation")
        print("Install with: pip install python-docx")

if __name__ == '__main__':
    main()
